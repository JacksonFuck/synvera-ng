"""Regression: .docx must NOT be routed to PDF parsers (they raise CalledProcessError).
Mirrors test_text_parser.py — DocxParser reads the document natively via python-docx."""
import docx as _docx

from raggw.parsing.docx_parser import DocxParser
from raggw.parsing.router import parse_file


def _make_docx(path):
    d = _docx.Document()
    d.add_heading("Protocolo POCUS", level=1)
    d.add_paragraph("Parágrafo um sobre ultrassom pulmonar.")
    d.add_paragraph("Parágrafo dois — linhas B e consolidação.")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Achado"
    t.rows[0].cells[1].text = "Linhas B difusas"
    d.save(str(path))


def test_docx_routes_to_docx_parser(tmp_path):
    f = tmp_path / "protocolo.docx"
    _make_docx(f)
    doc = parse_file(str(f))
    assert doc.parser == "docx"
    assert "ultrassom pulmonar" in doc.markdown


def test_docx_parser_extracts_headings_and_tables(tmp_path):
    f = tmp_path / "protocolo.docx"
    _make_docx(f)
    doc = DocxParser().parse(str(f))
    assert any(b.block_type == "heading" and "Protocolo POCUS" in b.text for b in doc.blocks)
    assert "Linhas B difusas" in doc.markdown
